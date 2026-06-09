"""
logic.py  –  Dataudtræk og Excel-skrivning til Excel Import-værktøjet

Tilpasset fra den originale extract_data.py.

Understøtter:
  - .docx-filer (fuldt implementeret)
  - .pdf-filer  (placeholder – klar til at implementere)

Tilføj ny filtype:
  1. Lav en extract_<type>(path) funktion der returnerer list[dict]
  2. Tilføj et elif-gren i extract_file()
"""

import os
import re
import shutil
import zipfile
import unicodedata
from difflib import SequenceMatcher
from datetime import date, datetime

import docx
import openpyxl
from openpyxl.worksheet.worksheet import Worksheet

# ─── Konstanter ───────────────────────────────────────────────────────────────

MONTH_MAP = {
    "januar": 1,   "februar": 2,  "marts": 3,    "april": 4,
    "maj": 5,      "juni": 6,     "juli": 7,     "august": 8,
    "september": 9, "oktober": 10, "november": 11, "december": 12,
}

EXPECTED_DATA_HEADERS = [
    "Dato", "Skole", "Skoletype", "Klassetrin", "Elever", "Lærere",
    "Undervisningsforløb", "Rundvisning", "PEH", "Downloadet Materiale",
    "Specialklasse", "ULF", "Aarhus Kommune",
]

# Kilder (hentet fra offentlige lister):
# - aarhus.dk: "Folkeskolerne i Aarhus" (kommunale + privat/friskoler)
# - suppleret med velkendte gymnasier/uddannelsesinstitutioner i Aarhus
AARHUS_SCHOOL_ALIASES = {
    # Kommunale folkeskoler/specialskoler
    "Bakkegårdsskolen", "Bavnehøj Skole", "Beder Skole", "Elev Skole", "Ellevangskolen",
    "Elsted Skole", "Engdalskolen", "Frederiksbjerg Skole", "Gammelgaardsskolen", "Hasle Skole",
    "Holme Skole", "Højvangskolen", "Hårup Skole", "Kaløvigskolen", "Katrinebjergskolen",
    "Kløverskolen", "Kragelundskolen", "Langagerskolen", "Lisbjergskolen", "Lystrup Skole",
    "Læssøesgades Skole", "Malling Skole", "Møllevangskolen", "Mårslet Skole", "Næshøjskolen",
    "Risskov Skole", "Rosenvangskolen", "Rundhøjskolen", "Sabro-Korsvejskolen", "Samsøgades Skole",
    "Skjoldhøjskolen", "Skovvangskolen", "Skæring Skole", "Skødstrup Skole", "Skåde Skole",
    "Solbjergskolen", "Stensagerskolen", "Strandskolen", "Sødalskolen", "Sølystskolen",
    "Søndervangskolen", "Tilst Skole", "Tranbjergskolen", "Vestergårdsskolen", "Viby Skole",
    "Virupskolen", "Vorrevangskolen", "Åby Skole", "Fokusskolen", "Netværksskolen",

    # Privat-/friskoler i Aarhus
    "Byskolen Gnist", "Børnenes Friskole", "Den Moderne Kulturelle Skole", "Elise Smiths Skole",
    "Egebakkeskolen", "Forældreskolen i Aarhus", "Friskolen Vildskud", "Højbjerg Privatskole",
    "Jakobskolen", "Laursens Realskole", "Leonardoskolen Aarhus", "N. Kochs Skole", "Interskolen",
    "Rudolf Steiner-Skolen", "Selam Friskole", "Skt. Knuds Skole", "Svalekærskolen",
    "Aarhus Friskole", "Aarhus Privatskole",

    # Ungdomsuddannelser/øvrige skoler i Aarhus (supplerende)
    "Aarhus Katedralskole", "Aarhus Statsgymnasium", "Århus Statsgymnasium", "Marselisborg Gymnasium",
    "Risskov Gymnasium", "Viby Gymnasium", "Egaa Gymnasium", "Egå Gymnasium", "Aarhus Gymnasium",
    "Langkær Gymnasium", "Aarhus Gymnasium Tilst", "Aarhus Tech", "Aarhus International School",
    "SOSU Østjylland", "SOSU Oestjylland",

    # --- Fra eksisterende Excel-ark (automatisk udtrukket) ---
    'A2B',
    'A2B Sprogcenter',
    'AABC Aarhus',
    'AABC HHX Gymnasiet Viby',
    'ADF Job og dansk',
    'AOF Job og Udd.',
    'AOF Sprogcenter',
    'AOF Sprogskole',
    'AOF, Job og dansk',
    'Aaby Fritidsklub',
    'Aakjærskolen',
    'Aarhus Akademi',
    'Aarhus Business College',
    'Aarhus Business School',
    'Aarhus Børnehøjskole',
    'Aarhus Designskole',
    'Aarhus Gym C / HTX',
    'Aarhus HF & VUC',
    'Aarhus HF og VUC',
    'Aarhus Handelsgymnasium',
    'Aarhus Handelsgymnasium, Viby',
    'Aarhus Handelskole',
    'Aarhus Lilleskole',
    'Aarhus Maskinmesterskole',
    'Aarhus Produktionsskole',
    'Aarhus STU',
    'Aarhus Statsskole',
    'Aarhus Tekniske Gymnasium',
    'Aarhus Universitet',
    'Aarhus Universitet - Institut for Kommunikation og Kultur',
    'Aarhus Universitet Kinastudier',
    'Aarhus Universitetshospital, psykiatrien',
    'Aarhus University, Department of Computer Science',
    'Alléskolen',
    'Alohe Skole',
    'Alternative Læringsarenaer (UngAarhus)',
    'Arkitektskolen i Aarhus',
    'Arveprinsesse Carolines Børneasyl',
    'Arveprinsesse Carolines Børnehave',
    'Arveprinsessen (bh)',
    'Arveprinsessens Børnehave',
    'Aspit Østjylland',
    'Asprit Østjylland',
    'BBH Gudrunsvej',
    'BHU Labyrinten',
    'BHU Thorshavnsgade',
    'BUA Psykiatri Skolebørn AUH',
    'Bakkegården',
    'Bakkely Privatskole',
    'Bakkeskolen Tilst',
    'Bakketoppen',
    'Bavnebakken',
    'Beder Kirke',
    'Bevægelserbørnhuset',
    'BevægelsesBørnehuset',
    'Bevægelseshuset',
    'Bh. Thorshavnsgade',
    'Bieskolen',
    'Birkehuset i Ry',
    'Birkeskolen',
    'Birkeskolen, STU',
    'Bo/Skole/Job',
    'Borgerskolen',
    'Brabrand Dagtilbud',
    'Brabrand Skole',
    'Brabrand Skole -  børnehuset Ilden',
    'Brovandeskolen',
    'Bulderby',
    'Bygningskonstruktørudd.',
    'Bøgeskov Skole',
    'Bøgeskovskolen',
    'Børne- og ungdomspsykiatrisk AUH',
    'Børnebåndet',
    'Børnebåndet Dagtilbud',
    'Børnegården Rundhøj',
    'Børnegården sct. Anna',
    'Børnehave',
    'Børnehaven Ahrendalsvej',
    'Børnehaven Arendalsvej',
    'Børnehaven Bjørnebakhus',
    'Børnehaven Engen',
    'Børnehaven Gammelgaarden',
    'Børnehaven Guldbryllupsasylet',
    'Børnehaven Hans Brogejvej',
    'Børnehaven Holkbækvej',
    'Børnehaven Jasminvej',
    'Børnehaven Labyrinten',
    'Børnehaven Lauge Koch',
    'Børnehaven Mælkevejen',
    'Børnehaven Ryhøjparken',
    'Børnehaven Rytterparken',
    'Børnehaven Skovtrolden',
    'Børnehaven Skt. Johannes',
    'Børnehaven Solnæs',
    'Børnehaven Svend Åge',
    'Børnehaven Thorshavnsgade',
    'Børnehaven Viben',
    'Børnehset Herluf Trolles Gade',
    'Børnehuset',
    'Børnehuset 8B',
    'Børnehuset Alsvej',
    'Børnehuset Bakken',
    'Børnehuset Besterbo',
    'Børnehuset Bylderup',
    'Børnehuset Bækken',
    'Børnehuset Børneliv',
    'Børnehuset Deruda',
    'Børnehuset Digterparken',
    'Børnehuset Egmontgården',
    'Børnehuset Elvehøjen',
    'Børnehuset Engblommevej',
    'Børnehuset Engen',
    'Børnehuset Fjældevænget',
    'Børnehuset Fjældvænget',
    'Børnehuset Frydenlund',
    'Børnehuset Fusijama',
    'Børnehuset Fusijana',
    'Børnehuset Gl. Ajstrup skole',
    'Børnehuset Haven',
    'Børnehuset Himmelblå',
    'Børnehuset Jernurten',
    'Børnehuset Jægergårdsgade',
    'Børnehuset Kilden',
    'Børnehuset Klokkervej',
    'Børnehuset Langenæs Allé',
    'Børnehuset Langenæsallé',
    'Børnehuset Luffen',
    'Børnehuset Luften',
    'Børnehuset Malmøgade',
    'Børnehuset Moltkesvej',
    'Børnehuset Mælkevejen',
    'Børnehuset Nordlys',
    'Børnehuset Ole Rømer',
    'Børnehuset Ovenpå',
    'Børnehuset Rybevænget',
    'Børnehuset Safiren',
    'Børnehuset Skoven',
    'Børnehuset Spiloppen',
    'Børnehuset Søsterhøj',
    'Børnehuset Thunfisken',
    'Børnehuset Tisviljana',
    'Børnehuset Trolden',
    'Børnehuset Troldeskoven',
    'Børnehuset Trolle',
    'Børnehuset Tusindfryd',
    'Børnehuset Valhalla',
    'Børnehuset Ved skoven',
    'Børnehuset Vestbyen',
    'Børnehuset Vesterbo',
    'Børnehuset Vores Sted',
    'Børnehuset bulerby',
    'Børnehuset rydevænget',
    'Børnehuset ved Åen',
    'Børnehøjen',
    'Børnekulturhuset',
    'Børnely',
    'Børnenes Have',
    'Castberggården',
    'Clavis Aarhus',
    'Clavis sprog- og kompetencecenter',
    'D.I.I. Regnbuen',
    'DII Anemonen',
    'DII Børnehuset Fjældevænget',
    'DII Børnehuset Troldevænget',
    'DII Børnenes Have',
    'DII Børneuniverset',
    'DII Delfinhuset',
    'DII Galaksen',
    'DII Hobbitien',
    'DII Hobbitten',
    'DII Hobitten',
    'DII Idrætsinst hoppelund',
    'DII Myretuen',
    'DII Regnbuen',
    'DII Rend og Hop',
    'DII SPILOPPEN HASSELAGER',
    'DII Skibet',
    'DII Skovkanten',
    'DII Solen',
    'DII Solsikken',
    'DII Solsikken Åbyhøj',
    'DII Træhøjen',
    'DII Træhøjen Viby',
    'DII Tumlehuset',
    'DII Tumlehøjen',
    'DII Venøvej',
    'DII Åtrolden',
    'DIN Friskole',
    'DNT Skole',
    'DTI Kompasset',
    'Daghuset Enggården',
    'Daginstitutionen Søskrænten',
    'Dagplejen',
    'Dagplejen Firkløveren',
    'Dagplejen Himmelblå',
    'Dagplejen Mårslet',
    'Dagplejen Skjoldhøj',
    'Dagplejen Skovvang',
    'Dagtilbud Nord',
    'Dagtilbud langenæsen',
    'Dagtilbuddet Børnehuset Vesterbo',
    'Dalgasskolen',
    'Den Grønne STU',
    'Den Private Børnepasning Klokkerbakken 56',
    'Den Private Pasningsordning Kastanjen',
    'Det Boligesociale Hus Gellerup',
    'Diakonhøjskolen',
    'Diakonhøjskolen Aarhus',
    'Dii Børnesnak',
    'Dr Alexandr Skole',
    'Dr. Alexandrines Børnehave',
    'Dr. Alexandrines Skole',
    'Dronning Alexandrines Børnehave',
    'Dybkær specialskole',
    'EUD Jordbrug Aarhus',
    'Efterskolen for scenekunst',
    'Egmentegården',
    'Egå Ungdomsskole',
    'Elkærskolen',
    'Ellehøjskolen',
    'Ellev Skoles Skole',
    'EngBlommen',
    'Engen',
    'Engen Flintebakken',
    'Engsøng Skole',
    'Erhvervsakademi Aarhus',
    'FGU Aarhus',
    'FGU Aarhus Skejby',
    'FGU Aarhus Vest',
    'FGU Egå',
    'FGU Himmerland',
    'FGU Skejby',
    'FGU Vest',
    'FGU Østjylland',
    'FK Gammelgård',
    'Fadls Børnehave',
    'Fadls Børnhave',
    'Familieforeningen Børneliv',
    'Fenrishus',
    'Fjældevænget',
    'Flintebakken',
    'Fonden Ørnehøj',
    'Frederiksbjerg Børnehus',
    'Frederiksbjerg Børnehuse',
    'Frederiksbjerg Klubben',
    'Frederikskirken',
    'Frederikskirken Skåde Sogn',
    'Frederiksskolen Skåde',
    'Fri læring',
    'Frisholm Skole',
    'Fritids og ungdomsklub Gammelgård',
    'Fritids- og ungdomsklubben Gammelgaard',
    'Fritidsklubben Gammelgaard',
    'Fritidsklubben Midtpunkt',
    'Frk. Ellen Gades Børnehave',
    'Fuglbjeggård',
    'Fuglebjerggaard',
    'GARDERØRET',
    'Gadekæret',
    'Galaksen',
    'Gammel Åby dagpleje',
    'Gilbroskolen',
    'Green Academy Aarhus',
    'Gunnergaardskolen',
    'Gyngehesten',
    'HALKÆRHØJSKOLEN I RISSKOV (AARHUS)',
    'HELTIDSWERKSKOLEN',
    'HHX Gymnasiet',
    'HHX Gymnasiet Risskov',
    'HHX Gymnasiet Viby',
    'HHX Gymnasiet i Risskov',
    'HHX Risskov',
    'Haldum-Hinnerup',
    'Handelsfagskolen Skåde',
    'Handelsgymnasiet Aarhus',
    'Hans Brages Parkens Børnehave',
    'Harlev Dagpleje',
    'Hasle Kirke',
    'Haven Vestervang',
    'Hjortshøj Kirke',
    'Hobditten',
    'Hojbjerg SFO',
    'Holmstrupgaard',
    'Holmstrupgård Interne Skole',
    'Hoppeland',
    'Hornslet Skole',
    'Hundslund Skole',
    'Hurlumhejhuset',
    'Hyldegårdsskolen',
    'Høj Skole',
    'Højbjerg Fritidsklub',
    'Høslund Efterskole',
    'Ida Holst Skole',
    'Idrætsbørnehuset Hoppeland',
    'Idrætsbørnehuset ved Globus1',
    'Idrætsinstitutionen Kolbøtten',
    'Idrætsinstitutionen Motalagade',
    'Idrætsinstitutionen Myrtuen',
    'Idrætsinstitutionen Rend og Hop',
    'Idrætsinstitutionen myretuen',
    'Idrætsinstitutonen Elkkær',
    'Incita Skole',
    'International Day Nursery Aarhud',
    'Johannes Voksne Center, Århus',
    'Jordbrugets Udannelsescenter',
    'Jordbrugets Udd.Center',
    'Jordbrugets UddannelsesCenter Århus',
    'Jordbrugets Uddannelsescenter',
    'Jordbrugets udd',
    'Junglen',
    'Juninakken, Skødstrup',
    'Kalengskolen',
    'Kareten',
    'Kareten Højbjerg',
    'Kinastudier, Aarhus Universitet',
    'Kirkevænget Skole',
    'Klaus Peters børnehave',
    'Klub Mega, østervangsskolen, hadsten',
    'Klub Perronen',
    'Klub Undergrund',
    'Klub2 Sydalsskolen',
    'Klubben Akva',
    'Klubben Barsøesgade',
    'Klubben Bodøgården',
    'Klubben Engdal',
    'Klubben Gammelgård',
    'Klubben Hjortshøj',
    'Klubben Holme',
    'Klubben Læssøesgade',
    'Klubben Nydam',
    'Klubben Peter Fabers Vej',
    'Klubben Peter Fabersvej',
    'Klubben Skødstrup',
    'Klubben Solbjerg',
    'Klubben Søndervang',
    'Klubben Trabjerg',
    'Klubben Tranbjerg',
    'Klubben Underground',
    'Klubberne Peter Fabers vej',
    'Klubhuset (UngiAarhus)',
    'Klubteket',
    'Kløverparkens Børnehave',
    'Knasten',
    'Kofoeds Skole',
    'Kolbøtten',
    'Kolt Klubben',
    'Konstruktørudd.',
    'Korshøjskolen',
    'Krabbehus Heldagsskole',
    'Kølbøtten',
    'Labyrinten',
    'Landsbyskolen',
    'Langenæs Dagtilbud',
    'Lauge Kochs Børnehave',
    'Lauge Kocks børnehave',
    'Lauge kocks børnhave',
    'Legehuset',
    'Legestuen Gyngehesten',
    'Leonardo Skolen',
    'Leonardskolen',
    'Lilleskolen - skoven',
    'Livøskolen',
    'Lundagerskolen (special)',
    'Lyngåskolen',
    'Lyngåskolen STU',
    'Læreruddannelsen Aarhus',
    'Læssøesgades klub',
    'Løgtengården',
    'MSP VIA University College',
    'Mammas House',
    'Mammas House, Beder',
    'Marielundskolen',
    'Mellemrummet',
    'Midtbyklubben',
    'Midtbyklubben Fritidsklub',
    'Multiplatform',
    'Musik- og Medieskolen Aarhus',
    'Myretuen',
    'Myretuen Ryhavevej',
    'Myrtuen Børnehave',
    'Mølholm skole',
    'Møllestien',
    'Natdrømmehaven i Langå',
    'Natur og Miljøenheden v. Århus',
    'Naturbørnehaven Mariendal',
    'Naturbørnehaven Viben',
    'Nordvang',
    'Ntur og Miljøjuset',
    'Ny tids læringsfællesskab',
    'Nørrestenbros Skovbørnehave',
    'PH Væksthuset',
    'PPO Klokkebakken',
    'Parat til start',
    'Peddersens Hus',
    'Privatskole',
    'Pædagogudd. Aarhus',
    'Pædagoguddannelsen Aarhus',
    'Pædagoguddannelsesen Aarhus',
    'Regnbuen',
    'Risskov Efterskole',
    'Rosenvangklubben',
    'Rundhøjsolen',
    'SFO Møllevang',
    'SFO Rudolf Steiner',
    'STRANDSØEN',
    'STU Hadsten',
    'STU Skolen',
    'Salem Privatskole',
    'Savstrup Dagtilbud',
    'Selam Privatskole',
    'Silkebølge',
    'Skattekisten',
    'Skelager Kirke',
    'Skjoldhøj DII',
    'Skjoldhøj Dagpleje',
    'Skolebyen',
    'Skolegades Børnehave',
    'Skolen for arkitektur',
    'Skolen ved Skarreso',
    'Skolevænget',
    'Skovby',
    'Skovby Børnehus',
    'Skovby Børnehus - Coldsmeden',
    'Skovby børnehus Guldsmeden',
    'Skovbørnehaven (Columbus)',
    'Skovgårdsparkens Børnehave',
    'Skovkanten',
    'Skrillingeskolen',
    'Skrivbyskolen',
    'Skt. Johannes Børnehave',
    'Skt. Johannes Børnesogn',
    'Skærbæk Distriktskole',
    'Skæring Egå Dagtilbud',
    'Sneglehuset',
    'Snurretoppen',
    'Snurretoppen tilst',
    'Social- og sundhedsskolen Aarhus',
    'Solbjerg Skoles SFO',
    'Solen',
    'Solhjem Rudolf Steiner',
    'Solsikken',
    'Solsikken DII',
    'Special Minds',
    'Spiloppen',
    'Springbræt',
    'Springbræt Aarhus',
    'Stationsvangen Skødstrup',
    'Steiner HF',
    'Steiner HF Aarhus',
    'Storkereden',
    'Studenterhus Aarhus',
    'SØVÆNGBSSKOLEN',
    'Søndermark Skolen',
    'Søndervang Vestergård Dagtilbud',
    'Søndervang dagtilbud',
    'Søndervangs Dagtilbud',
    'Søskrænten',
    'TCU hvinningdalskolen',
    'Tarikvejeus skole',
    'Thunfisken',
    'Thurshavnsgade',
    'Tilst SFO',
    'Topdalskolen',
    'Tranbjerg Fritidsklub',
    'Tranbjerg Skoles SFO',
    'Træhøjen Søndervang Vestergård',
    'Tumlehuset',
    'Tumlehøjen',
    'UC Plus Højbjerg',
    'ULF i Aarhus',
    'Uddannelsesinstitutionen Facet',
    'Ungdomsskolen',
    'Ungecenter USF',
    'Ungesbjerg',
    'UngiAarhus',
    'UngiÅrhus',
    'Universet, Kolt Hasselager',
    'VIA',
    'VIA Aarhus pæd.',
    'VIA Campus C',
    'VIA Film',
    'VIA International',
    'VIA Lærerudd.',
    'VIA Pæd. Aarhus',
    'VIA Pædagogudd.',
    'VIA UC',
    'VIA University Colleg',
    'VIA University College - Campus C',
    'VIA University College Aarhus',
    'VIA University College Lærerudd.',
    'VIA University College, pæd.',
    'Veng Skole Idrætsbør',
    'Vestfjendesskolen',
    'Vestskolen afd. Skovbakken',
    'Via University College',
    'Viben',
    'Vidtskue',
    'Virring skole',
    'Vokseværket',
    'Vuggestuen Gyrvelhaven',
    'Vuggestuen Gårdhaven',
    'Vuggestuen Kernehuset',
    'Vuggestuen Kernehuset (Vestereng 12)',
    'Vuggestuen Kornbakken',
    'Vuggestuen Langenæsstien',
    'Vuggestuen Myretuen',
    'Vuggestuen Nannasvej',
    'Vuggestuen Solsikken',
    'Vuggestuen Solskikken',
    'Vuggestuen ved æbletræet',
    'Væksthuset',
    'Væksthuset 8210',
    'Væksthuset skole',
    'al hikma skole',
    'birkehuset',
    'firkløverskolen afd elkjær',
    'jelling friskole',
    'learnmark',
    'natur -og miljø børnehuset viden',
    'natur og miljøbørnehuset Viben',
    'nørreå friskole',
    'rønde skole',
    'skt janus børnehave',
    'stenagergård',
    'stilleng skole',
    'toflelunden, engen',
    'tungelundskolen',
    'vestervang skole',
    'Åby Børnehave',
    'Åby Dagpleje',
    'Åby Fritids- og Ungdomsklub',
    'Åby Fritidsklub',
    'Åby Sogn',
    'Århus STU',
    'Århus STU Syd',
    'Århus Universitet',
    'Åskolen',
    'Åtrolden',
    'Østerbo Vuggestue',
    'Østre Skole Middelfart',
    'østerbyskolen',
    # --- Officielle Aarhus Kommune dagtilbud og afdelinger (aarhus.dk) ---
    '2-Kløveren',
    'Beder-Malling Dagtilbud',
    'Brabrandsøens Dagtilbud',
    'Børnegården Bifrost',
    'Børneslottet',
    'Børneslottet Bavnehøj',
    'Børnesymfonien',
    'DII Ellevang',
    'DII Tornebakken',
    'Dagtilbud Hasle',
    'Dagtilbuddet Aarhus Ø',
    'Dagtilbuddet Christiansbjerg',
    'Damhuset',
    'Det Blå Hus',
    'Ellevang Dagtilbud',
    'Elverhøj',
    'Eventyrskoven',
    'Flyverhøjen',
    'Frederiksbjerg Dagtilbud',
    'Gl. Åby Dagtilbud',
    'Grøften',
    'Harlev Dagtilbud',
    'Holme Børnehus',
    'Holme-Rundhøj Dagtilbud',
    'Holmetræet',
    'Hårup Børnegård',
    'Idrætsinstitutionen Ellekær',
    'International Day Nursery Aarhus',
    'Kløverhulen',
    'Kolt Hasselager Dagtilbud',
    'Kongevellen',
    'Kridthuset',
    'Kroghsgade Vuggestue',
    'Krokodillehaven',
    'Kræmmerhuset',
    'Lisbjerg-Trige-Spørring Dagtilbud',
    'Livstræet',
    'Lystrup-Elsted Dagtilbud',
    'Mariehønen',
    'Midtbyens Dagtilbud',
    'Møllehuset',
    'Møllevang Dagtilbud',
    'Nordbyernes Dagtilbud',
    'Pilehytten',
    'Risskov Børnehus',
    'Risskov Dagtilbud',
    'Sabro Dagtilbud',
    'Skejby Vorrevang Dagtilbud',
    'Skjoldhøj Dagtilbud',
    'Skovbørnehaven Grøftekanten',
    'Skovvang og Trøjborg Dagtilbud',
    'Skuden Skram',
    'Skåde-Højbjerg Dagtilbud',
    'Skæring-Sølyst Dagtilbud',
    'Skødstrup Dagtilbud',
    'Solbjerg-Mårslet Dagtilbud',
    'Specialdagtilbud Skovbrynet',
    'Spiren',
    'Spirrevippen',
    'Stavtrup Dagtilbud',
    'Stennehøj Allé',
    'Strandens Dagtilbud',
    'Søhøjen',
    'Tilst Dagtilbud',
    'Todbjerg Børnehus',
    'Tovshøj Dagtilbud',
    'Tranbjerg Dagtilbud',
    'Troldehøjen',
    'Trøjborg Børnehus',
    'Under Bøgen',
    'Valhalla',
    'Venøvej 9',
    'Viby-Rosenvang Dagtilbud',
    'Vuggestuen Bjørnbaksvej',
    'Vuggestuen Charlottehøj',
    'Vuggestuen Forteleddet',
    'Vuggestuen Graven',
    'Vuggestuen Hasle',
    'Vuggestuen Jægergårdsgade',
    'Vuggestuen Krible Krable',
    'Vuggestuen Lindenborgvej',
    'Vuggestuen Marselis Boulevard',
    'Vuggestuen Ny Vimmerby',
    'Vuggestuen Rudgården',
    'Vuggestuen Solstrålen',
    'Vuggestuen Trekanten',
    'Vuggestuen Trøjborg',
    'Vuggestuen Vestervang',
    'Vuggestuen Århusbo',
    'Åbyhøj Dagtilbud',
    'Æsken',
    # --- Officielle Aarhus Kommune dagtilbud og afdelinger (aarhus.dk) ---
    'Blomsterhaven',
    'BørneUniverset',
    'Børnesnak',
    'DII Annekæret',
    'DII Evigglad',
    'DII Gl. Egå Børnehus',
    'DII Humlebien',
    'DII Kompasset',
    'DII Skærgården',
    'DII Skæringhus',
    'DII Sommerfuglen',
    'DII Stenhøjen',
    'DII Studsen',
    'DII Tryllefløjten',
    'Den Integrerede Institution Drivhuset',
    'Den Integrerede Institution Haven',
    'Elsted Børnehus',
    'Engskovgård',
    'Fuglbjerggård',
    'Hasselhaven',
    'Hasselhuset',
    'Hobbitten',
    'Huset på Bakken',
    'Junibakken',
    'Klokkeskoven',
    'Kometen',
    'Kærgården',
    'Langenæsstien 6A',
    'Langenæsstien 6B',
    'Mårslet Børneunivers',
    'Mårslet Skovbørnehus',
    'Naturbørnehaven Ajstrup Gl. Skole',
    'Nymarken',
    'Paletten',
    'Perlen',
    'Satellitten Lisbjerg',
    'Skibet',
    'Solskin',
    'Spørring Børnehus',
    'Søstjernen',
    'Vuggestuen Ankersgade',
    'Vuggestuen Brohaven',
    'Vuggestuen Den Grønne Gren',
    'Vuggestuen Firkløveren',
    'Vuggestuen Heimdalsvej',
    'Vuggestuen Kløvervangen',
    'Vuggestuen Lillefod',
    'Vuggestuen Lærkebo',
    'Vuggestuen Læssøesgade',
    'Vuggestuen Pilehuset',
    'Vuggestuen Silkeborgvej',
    'Vuggestuen Tommelise',
    'Vuggestuen Tusindfryd',
    'Vuggestuen Åhøjen',
    'Vuggestuen Åkanden',
    'Ørnereden',
}


def _normalize_school_name(value: str | None) -> str:
    """Normaliserer skolenavne for robust sammenligning."""
    text = "" if value is None else str(value)
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower().replace("-", " ")
    text = re.sub(r"[^a-z0-9 ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


_AARHUS_ALIAS_NORMALIZED = {
    _normalize_school_name(name) for name in AARHUS_SCHOOL_ALIASES
}


def create_excel_backup(excel_path: str, reason: str = "update") -> str:
    """
    Opretter tidsstemplet backup af en Excel-fil i en lokal backup-mappe.
    Returnerer stien til backup-filen.
    """
    source = os.path.abspath(excel_path)
    if not os.path.exists(source):
        raise FileNotFoundError(f"Excel-fil findes ikke: {source}")

    folder = os.path.dirname(source)
    backup_dir = os.path.join(folder, "_backups")
    os.makedirs(backup_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(source))[0]
    safe_reason = re.sub(r"[^a-zA-Z0-9_-]+", "_", reason).strip("_") or "update"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    candidate = os.path.join(backup_dir, f"{base_name}__{stamp}__{safe_reason}.xlsx")
    counter = 1
    while os.path.exists(candidate):
        candidate = os.path.join(
            backup_dir,
            f"{base_name}__{stamp}__{safe_reason}_{counter}.xlsx",
        )
        counter += 1

    shutil.copy2(source, candidate)
    return candidate


def _is_aarhus_school(name: str | None) -> bool:
    """Matcher indkommende skolenavn mod kendt Aarhus-skoleliste."""
    candidate = _normalize_school_name(name)
    if not candidate:
        return False

    if candidate in _AARHUS_ALIAS_NORMALIZED:
        return True

    # Tillad små variationer, fx "Sølystskolen Egå" eller "Aarhus Gymnasium, Tilst"
    for alias in _AARHUS_ALIAS_NORMALIZED:
        if alias and (candidate.startswith(alias) or alias.startswith(candidate) or f" {alias} " in f" {candidate} "):
            return True

    # Let fuzzy-match ved mindre stavefejl/forkortelser.
    # Eksempel: "Aarhus Gymnasiun" eller "Aarhus Gym".
    school_markers = (
        "skol", "gymnasi", "katedralskol", "friskol",
        "privatskol", "tech", "sosu", "stx", "hf"
    )
    if not any(marker in candidate for marker in school_markers):
        return False

    best_ratio = 0.0
    for alias in _AARHUS_ALIAS_NORMALIZED:
        ratio = SequenceMatcher(None, candidate, alias).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            if best_ratio >= 0.84:
                return True

    return False


# ─── Dato-parsing ──────────────────────────────────────────────────────────────

def parse_date(filename: str) -> date | None:
    """Udtrækker dato fra filnavne som '12. Januar - ...' eller '15 april - ...'"""
    name = os.path.splitext(os.path.basename(filename))[0]
    match = re.search(r"(\d{1,2})[.\s]+([a-zA-ZæøåÆØÅ]+)", name)
    if match:
        day = int(match.group(1))
        month = MONTH_MAP.get(match.group(2).lower())
        if month:
            return date(2026, month, day)
    return None


# ─── .docx-hjælpefunktioner ────────────────────────────────────────────────────

def _resolve_dittos(rows: list[list[str]]) -> list[list[str]]:
    """Erstatter ditto-tegn ('"') med værdien fra forrige datarække."""
    prev = [None] * 8
    result = []
    for row in rows:
        resolved = []
        for i, cell in enumerate(row):
            if str(cell).strip() in {'"', '”', '“'} and prev[i] is not None:
                resolved.append(prev[i])
            else:
                resolved.append(cell)
        for i, cell in enumerate(resolved):
            if cell:
                prev[i] = cell
        result.append(resolved)
    return result


def _section_type(header: list[str]) -> str:
    """Bestemmer sektions-label ud fra header-tekst."""
    h = header[1].upper()
    if "FORLØB" in h or "RUNDVISNING" in h:
        return "Forløb/Rundvisning"
    if "SPECIAL" in h or "BØRNEHAVE" in h:
        return "Special/Børnehave"
    return "PEH"


def _extract_docx_sections(table) -> list[tuple]:
    """Opdeler tabel-rækker i sektioner baseret på '#'-header."""
    rows = [[c.text.strip() for c in row.cells] for row in table.rows]
    starts = [i for i, r in enumerate(rows) if r[0] == "#"]
    sections = []
    for idx, start in enumerate(starts):
        end = starts[idx + 1] if idx + 1 < len(starts) else len(rows)
        data = [r for r in rows[start + 1:end] if r[0].upper() != "EX" and any(r[1:])]
        sections.append((rows[start], data))
    return sections


# ─── Filtype-udtrækkere ────────────────────────────────────────────────────────

def extract_docx(path: str) -> list[dict]:
    """Udtrækker datarækker fra en .docx-fil."""
    doc_date = parse_date(path)
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"Kunne ikke åbne {os.path.basename(path)}: {e}") from e

    if not doc.tables:
        return []

    collected = []
    for header, data_rows in _extract_docx_sections(doc.tables[0]):
        label = _section_type(header)
        for row in _resolve_dittos(data_rows):
            collected.append({
                "Dato":           doc_date,
                "Type":           label,
                "#":              row[0],
                "Navn":           row[1],
                "Klasse":         row[2],
                "Antal":          row[3],
                "Lærer/Kontakt":  row[4],
                "Kontakt/Ordre":  row[5],
                "Tid":            row[6],
                "Initialer":      row[7] if len(row) > 7 else "",
            })
    return collected


def extract_pdf(path: str) -> list[dict]:
    """
    PDF-udtræk – endnu ikke implementeret.
    Implementér her når PDF-support ønskes.
    """
    raise NotImplementedError(
        f"PDF-import er ikke implementeret endnu: {os.path.basename(path)}"
    )


def _normalize_header(value) -> str:
    """Normaliserer headertekst så små stavevariationer ikke vælter sammenligningen."""
    text = "" if value is None else str(value)
    text = text.strip().lower()
    return "".join(ch for ch in text if ch.isalnum())


def _repair_missing_pivot_cache_records(
    excel_path: str,
    create_backup_before_change: bool = True,
) -> list[str]:
    """
    Nogle Excel-filer har pivot-cache relationer til records-filer som mangler.
    openpyxl fejler på disse filer. Denne funktion opretter minimale records-filer
    så workbooken kan åbnes og skrives igen.

    Returnerer en liste af oprettede filstier inde i xlsx-zip'en.
    """
    with zipfile.ZipFile(excel_path, "r") as z:
        names = set(z.namelist())
        rel_files = [
            n for n in names
            if n.startswith("xl/pivotCache/_rels/") and n.endswith(".rels")
        ]

        needed: set[str] = set()
        for rel in rel_files:
            rel_xml = z.read(rel).decode("utf-8", errors="ignore")
            for target in re.findall(r'Target="([^"]+)"', rel_xml):
                normalized = target.lstrip("/")
                if "pivotCacheRecords" in normalized:
                    needed.add(normalized)

    missing = sorted(p for p in needed if p not in names)
    if not missing:
        return []

    if create_backup_before_change:
        create_excel_backup(excel_path, reason="pivot_repair")

    minimal_records = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<pivotCacheRecords xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="0"/>'
    )

    tmp_path = excel_path + ".tmp"
    with zipfile.ZipFile(excel_path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, zin.read(item.filename))
            for target in missing:
                zout.writestr(target, minimal_records)

    os.replace(tmp_path, excel_path)
    return missing


def _sheet_headers(ws: Worksheet, width: int = 13) -> list[str]:
    """Henter første række i arket (de første `width` kolonner)."""
    return [ws.cell(row=1, column=i).value for i in range(1, width + 1)]


def is_data_sheet(ws: Worksheet) -> bool:
    """
    Returnerer True hvis arket matcher den forventede dataskabelon.
    Pivot-/opsummeringsark returnerer False.
    """
    actual = [_normalize_header(h) for h in _sheet_headers(ws, len(EXPECTED_DATA_HEADERS))]
    expected = [_normalize_header(h) for h in EXPECTED_DATA_HEADERS]
    return actual == expected


def _xml_cell_value(cell_el: str, shared_strings: list[str]) -> str:
    """
    Udtrækker tekst fra et <c>-element givet den rå XML og shared-strings-listen.
    Understøtter shared strings (t="s"), inline strings (t="inlineStr") og tal.
    """
    t_attr = re.search(r'\bt="(\w+)"', cell_el)
    t = t_attr.group(1) if t_attr else None
    v = re.search(r'<v>([^<]*)</v>', cell_el)
    if t == "s" and v and shared_strings:
        try:
            return shared_strings[int(v.group(1))]
        except (IndexError, ValueError):
            return ""
    if t == "inlineStr":
        m = re.search(r'<t[^>]*>([^<]*)</t>', cell_el)
        return m.group(1) if m else ""
    return v.group(1) if v else ""


def _fast_sheet_headers(sheet_xml: str, shared_strings: list[str]) -> list[str]:
    """Læser første datarække (r=1) fra sheet XML og returnerer cellernes værdier."""
    row1 = re.search(r'<row\b[^>]*\br="1"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
    if not row1:
        return []
    cells = re.findall(r'<c\b[^>]*>.*?</c>', row1.group(1), re.DOTALL)
    result = []
    for cell in cells:
        ref = re.search(r'\br="([A-Z]+)\d+"', cell)
        col_letter = ref.group(1) if ref else None
        if col_letter:
            col_idx = 0
            for ch in col_letter:
                col_idx = col_idx * 26 + (ord(ch) - ord('A') + 1)
            # Fyld huller med tomme strenge
            while len(result) < col_idx - 1:
                result.append("")
            result.append(_xml_cell_value(cell, shared_strings))
    return result


def _load_shared_strings(z: zipfile.ZipFile) -> list[str]:
    """Indlæser sharedStrings.xml og returnerer en liste af strengværdier."""
    if "xl/sharedStrings.xml" not in z.namelist():
        return []
    xml = z.read("xl/sharedStrings.xml").decode("utf-8", errors="replace")
    # Hent alle <si>-elementer og udtræk tekst
    result = []
    for si in re.findall(r'<si>(.*?)</si>', xml, re.DOTALL):
        texts = re.findall(r'<t[^>]*>([^<]*)</t>', si)
        result.append("".join(texts))
    return result


def list_importable_sheets(excel_path: str) -> list[str]:
    """
    Finder alle ark i en Excel-fil som matcher datastrukturen.
    Bruger direkte XML-læsning (ingen openpyxl) for at undgå timeout
    på komplekse filer med mange pivot-tabeller.
    """
    expected = [_normalize_header(h) for h in EXPECTED_DATA_HEADERS]
    try:
        with zipfile.ZipFile(excel_path, "r") as z:
            shared_strings = _load_shared_strings(z)
            wb_xml = z.read("xl/workbook.xml").decode("utf-8", errors="ignore")
            rels_xml = z.read("xl/_rels/workbook.xml.rels").decode("utf-8", errors="ignore")

            # Byg rId -> Target map
            rid_to_target: dict[str, str] = {}
            for m in re.finditer(r'<Relationship[^>]+Id="([^"]+)"[^>]+Target="([^"]+)"', rels_xml):
                rid_to_target[m.group(1)] = m.group(2)

            result = []
            for m in re.finditer(r'<sheet\b([^/]*/?>)', wb_xml):
                attrs = m.group(1)
                name_m = re.search(r'\bname="([^"]+)"', attrs)
                rid_m  = re.search(r'\br:id="([^"]+)"', attrs)
                state_m = re.search(r'\bstate="([^"]+)"', attrs)
                if not name_m or not rid_m:
                    continue
                if state_m and state_m.group(1) != "visible":
                    continue

                sheet_name = name_m.group(1)
                rid = rid_m.group(1)
                target = rid_to_target.get(rid, "")
                zip_path = "xl/" + target if not target.startswith("xl/") else target

                if zip_path not in z.namelist():
                    continue

                sheet_xml = z.read(zip_path).decode("utf-8", errors="replace")
                headers = [_normalize_header(h) for h in _fast_sheet_headers(sheet_xml, shared_strings)]
                if headers[:len(expected)] == expected:
                    result.append(sheet_name)

        return result
    except Exception as e:
        raise ValueError(f"Kunne ikke læse Excel-filen: {e}") from e


def _parse_antal(antal_raw: str) -> tuple[object, object]:
    """
    Parser feltet 'Antal' til (elever, lærere).
    Eksempler:
      '25+2' -> (25, 2)
      '28'   -> (28, None)
      '12-16' eller ukendt tekst -> ('12-16', None)
    """
    raw = "" if antal_raw is None else str(antal_raw).strip()
    if not raw:
        return None, None

    compact = raw.replace(" ", "")
    if "+" in compact:
        left, right = compact.split("+", 1)
        if left.isdigit() and right.isdigit():
            return int(left), int(right)

    if compact.isdigit():
        return int(compact), None

    return raw, None


def _to_year_row(row: dict) -> dict:
    """Mapper én udtrukket række til kolonnerne i årsarket (fx 2026)."""
    elever, laerere = _parse_antal(row.get("Antal"))

    result = {
        "Dato": row.get("Dato"),
        "Skole": row.get("Navn"),
        "Skoletype": None,
        "Klassetrin": row.get("Klasse"),
        "Elever": elever,
        "Lærere": laerere,
        "Undervisningsforløb": None,
        "Rundvisning": None,
        "PEH": None,
        "Downloadet Materiale": None,
        "Specialklasse": None,
        "ULF": None,
        "Aarhus Kommune": "x" if _is_aarhus_school(row.get("Navn")) else None,
    }

    section_type = row.get("Type")
    if section_type == "PEH":
        result["PEH"] = "x"
    elif section_type == "Special/Børnehave":
        result["Specialklasse"] = "x"
    elif section_type == "Forløb/Rundvisning":
        # Kildedata skelner ikke entydigt mellem disse to i nuværende parser.
        result["Undervisningsforløb"] = "x"

    return result


# ─── Dispatcher ───────────────────────────────────────────────────────────────

def extract_file(path: str) -> tuple[list[dict], str | None]:
    """
    Udtrækker data fra én fil.
    Returnerer (rækker, fejlbesked).
    fejlbesked er None ved succes.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            return extract_docx(path), None
        elif ext == ".pdf":
            return extract_pdf(path), None
        else:
            return [], f"Ukendt filtype: {ext}"
    except NotImplementedError as e:
        return [], str(e)
    except Exception as e:
        return [], str(e)


# ─── Excel-skriver ────────────────────────────────────────────────────────────

def _excel_serial_date(d) -> int:
    """Konverterer en date/datetime til Excel's serielle datotal (1900-system)."""
    if isinstance(d, datetime):
        d = d.date()
    return (d - date(1899, 12, 30)).days


def _xml_escape(text: str) -> str:
    """Escaper tekst til brug i XML."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


def _validate_data_sheet(excel_path: str, sheet_name: str) -> None:
    """Tjekker at arket findes og matcher datastrukturen. Kaster ValueError ellers."""
    sheets = list_importable_sheets(excel_path)
    if sheet_name not in [s for s in _all_sheet_names(excel_path)]:
        raise ValueError(f"Ark '{sheet_name}' findes ikke i Excel-filen.")
    if sheet_name not in sheets:
        raise ValueError(
            f"Ark '{sheet_name}' er ikke et data-ark med korrekt struktur. "
            "Vælg et årsark (fx 2026), ikke et pivot-/Data-ark."
        )


def _all_sheet_names(excel_path: str) -> list[str]:
    """Returnerer alle ark-navne (visible + hidden) uden at åbne med openpyxl."""
    with zipfile.ZipFile(excel_path, "r") as z:
        wb_xml = z.read("xl/workbook.xml").decode("utf-8", errors="ignore")
    return re.findall(r'<sheet\b[^>]*\bname="([^"]+)"', wb_xml)


def write_to_excel(
    rows: list[dict],
    excel_path: str,
    sheet_name: str,
    progress_callback=None,
) -> tuple[int, list[str]]:
    """
    Tilføjer rækker til et eksisterende data-ark i en Excel-fil.

    Bruger ren XML-kirurgi: arkets XML læses direkte fra zip'en, nye rækker
    tilføjes med inline-strenge, og kun det ændrede ark + tabel + pivot-cache
    skrives om. Resten af filen (kolonnebredder, pivot-tabeller, diagrammer,
    formatering) bevares byte-for-byte. openpyxl bruges KUN til validering –
    aldrig til at gemme – fordi en openpyxl-roundtrip ødelægger komplekse filer.

    Returns:
        (antal_rækker_skrevet, liste_af_advarsler)
    """
    warnings: list[str] = []

    _validate_data_sheet(excel_path, sheet_name)

    mapped_rows = [_to_year_row(r) for r in rows]
    total = len(mapped_rows)
    if total == 0:
        return 0, warnings

    sheet_zip_path = _find_sheet_zip_path(excel_path, sheet_name)
    if not sheet_zip_path:
        raise ValueError(f"Kunne ikke finde ark-filen for '{sheet_name}'.")

    backup_path = create_excel_backup(excel_path, reason="excel_import")
    warnings.append(f"Backup oprettet: {os.path.basename(backup_path)}")

    # Læs arkets XML direkte fra zip'en
    with zipfile.ZipFile(excel_path, "r") as z:
        sheet_xml = z.read(sheet_zip_path).decode("utf-8", errors="replace")

    # Find højeste eksisterende rækkenummer
    existing_rows = [int(m) for m in re.findall(r'<row r="(\d+)"', sheet_xml)]
    last_existing = max(existing_rows) if existing_rows else 1
    start_row = last_existing + 1

    # Find den style der bruges til dato-kolonnen (A) i eksisterende rækker
    a_styles = re.findall(r'<c r="A\d+"[^>]*\bs="(\d+)"', sheet_xml)
    date_style = a_styles[-1] if a_styles else "6"

    # Kolonner: (bogstav, nøgle, er_tal_kolonne)
    columns = [
        ("A", "Dato", False),
        ("B", "Skole", False),
        ("C", "Skoletype", False),
        ("D", "Klassetrin", True),
        ("E", "Elever", True),
        ("F", "Lærere", True),
        ("G", "Undervisningsforløb", False),
        ("H", "Rundvisning", False),
        ("I", "PEH", False),
        ("J", "Downloadet Materiale", False),
        ("K", "Specialklasse", False),
        ("L", "ULF", False),
        ("M", "Aarhus Kommune", False),
    ]

    def _is_number(v) -> bool:
        return isinstance(v, (int, float)) and not isinstance(v, bool)

    new_rows_xml: list[str] = []
    for i, row_data in enumerate(mapped_rows):
        r = start_row + i
        cells: list[str] = []
        for col, key, is_num_col in columns:
            val = row_data.get(key)
            if val is None or val == "":
                continue

            if key == "Dato" and isinstance(val, (date, datetime)):
                serial = _excel_serial_date(val)
                cells.append(f'<c r="{col}{r}" s="{date_style}"><v>{serial}</v></c>')
            elif is_num_col and _is_number(val):
                cells.append(f'<c r="{col}{r}"><v>{val}</v></c>')
            else:
                text = _xml_escape(str(val))
                cells.append(
                    f'<c r="{col}{r}" t="inlineStr"><is>'
                    f'<t xml:space="preserve">{text}</t></is></c>'
                )

        new_rows_xml.append(
            f'<row r="{r}" spans="1:13" x14ac:dyDescent="0.35">{"".join(cells)}</row>'
        )

        if progress_callback:
            progress_callback((i + 1) / total, f"Skriver række {i + 1} / {total}")

    last_row = start_row + total - 1

    # Indsæt de nye rækker lige før </sheetData>
    insert_block = "".join(new_rows_xml)
    if "</sheetData>" in sheet_xml:
        sheet_xml = sheet_xml.replace("</sheetData>", insert_block + "</sheetData>", 1)
    else:
        raise ValueError("Ugyldig ark-XML: mangler </sheetData>.")

    # Opdater dimension-ref
    sheet_xml = re.sub(
        r'(<dimension ref=")[A-Z]+\d+:[A-Z]+\d+(")',
        rf'\g<1>A1:M{last_row}\g<2>',
        sheet_xml, count=1,
    )

    sheet_bytes = sheet_xml.encode("utf-8")

    # Patch tabel-ref og pivot-cache-refs direkte fra originalen
    xml_patches = _patch_xlsx_xml(excel_path, sheet_name, last_row)

    # Saml alle erstatninger: arkets nye XML + patchede tabel/pivot-filer
    replacements: dict[str, bytes] = {sheet_zip_path: sheet_bytes}
    replacements.update(xml_patches)

    # Byg den endelige fil: kopiér ALT fra originalen, erstat kun de ændrede.
    # Hver post skrives præcis én gang for at undgå dublet-korruption.
    tmp_final = excel_path + ".tmp"
    with zipfile.ZipFile(excel_path, "r") as z_orig:
        with zipfile.ZipFile(tmp_final, "w", zipfile.ZIP_DEFLATED) as z_out:
            for item in z_orig.infolist():
                if item.filename in replacements:
                    z_out.writestr(item, replacements[item.filename])
                else:
                    z_out.writestr(item, z_orig.read(item.filename))

    os.replace(tmp_final, excel_path)

    return total, warnings


def _patch_xlsx_xml(excel_path: str, sheet_name: str, new_last_row: int) -> dict[str, bytes]:
    """
    Læser og patcher XML-filer direkte fra den originale xlsx:
      - xl/tables/tableN.xml for det pågældende ark: opdaterer ref= på <table>-elementet
      - xl/pivotCache/pivotCacheDefinition*.xml: opdaterer worksheetSource ref= til A1:M{new_last_row}
    Returnerer en dict {intern_zip_sti: patched_bytes}.
    """
    patched: dict[str, bytes] = {}
    try:
        with zipfile.ZipFile(excel_path, "r") as z:
            # 1. Patch tabel-filer der tilhører dette ark
            sheet_zip_path = _find_sheet_zip_path(excel_path, sheet_name)
            if sheet_zip_path:
                sheet_rels = _sheet_rels_path(sheet_zip_path)
                if sheet_rels in z.namelist():
                    rels_xml = z.read(sheet_rels).decode("utf-8", errors="ignore")
                    for target in re.findall(r'Target="([^"]+)"', rels_xml):
                        if re.search(r"/tables/table\d+\.xml$", target, re.IGNORECASE):
                            sheet_dir = "/".join(sheet_zip_path.split("/")[:-1])
                            norm = os.path.normpath(sheet_dir + "/" + target).replace("\\", "/")
                            if norm in z.namelist():
                                xml = z.read(norm).decode("utf-8", errors="replace")
                                xml = re.sub(
                                    r'(<table\b[^>]*\bref=")[A-Z]+\d+:[A-Z]+\d+"',
                                    rf'\g<1>A1:M{new_last_row}"',
                                    xml, count=1,
                                )
                                patched[norm] = xml.encode("utf-8")

            # 2. Patch pivot-cache-definitioner
            for name in z.namelist():
                if re.search(r"xl/pivotCache/pivotCacheDefinition\d*\.xml$", name):
                    xml = z.read(name).decode("utf-8", errors="replace")

                    def _patch_source(m: re.Match) -> str:
                        tag = m.group(0)
                        if f'sheet="{sheet_name}"' in tag or f"sheet='{sheet_name}'" in tag:
                            tag = re.sub(r'ref="[A-Z]+\d+:[A-Z]+\d+"', f'ref="A1:M{new_last_row}"', tag)
                        return tag

                    xml = re.sub(r"<worksheetSource\b[^>]*/?>", _patch_source, xml)
                    patched[name] = xml.encode("utf-8")
    except Exception:
        pass
    return patched


def _sheet_rels_path(sheet_zip_path: str) -> str:
    """Konverterer 'xl/worksheets/sheet2.xml' til 'xl/worksheets/_rels/sheet2.xml.rels'."""
    parts = sheet_zip_path.rsplit("/", 1)
    return parts[0] + "/_rels/" + parts[1] + ".rels"


def _find_sheet_zip_path(excel_path: str, sheet_name: str) -> str | None:
    """
    Finder den interne zip-sti (fx 'xl/worksheets/sheet2.xml') for et ark
    med det givne navn ved at parse workbook.xml og workbook.xml.rels.
    """
    try:
        with zipfile.ZipFile(excel_path, "r") as z:
            wb_xml = z.read("xl/workbook.xml").decode("utf-8", errors="ignore")
            rels_xml = z.read("xl/_rels/workbook.xml.rels").decode("utf-8", errors="ignore")

        # Find rId for arket
        rid_match = re.search(
            r'<sheet\b[^>]*\bname="' + re.escape(sheet_name) + r'"[^>]*\br:id="([^"]+)"',
            wb_xml,
        )
        if not rid_match:
            rid_match = re.search(
                r'<sheet\b[^>]*\br:id="([^"]+)"[^>]*\bname="' + re.escape(sheet_name) + r'"',
                wb_xml,
            )
        if not rid_match:
            return None
        rid = rid_match.group(1)

        # Find fil-Target for rId'en
        target_match = re.search(
            r'<Relationship\b[^>]*\bId="' + re.escape(rid) + r'"[^>]*\bTarget="([^"]+)"',
            rels_xml,
        )
        if not target_match:
            target_match = re.search(
                r'<Relationship\b[^>]*\bTarget="([^"]+)"[^>]*\bId="' + re.escape(rid) + r'"',
                rels_xml,
            )
        if not target_match:
            return None

        target = target_match.group(1).lstrip("/")
        # Target er relativt til xl/, fx "worksheets/sheet2.xml"
        return "xl/" + target if not target.startswith("xl/") else target
    except Exception:
        return None
